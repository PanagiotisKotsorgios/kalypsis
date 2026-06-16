"""Generates a Greek presentation Word document for insurance agencies (γραφεία).

Run from the repo root:
    python docs/generate_brochure.py

Output: docs/Kalypsis_Platform_Greek.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Inches, Pt, RGBColor, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

ROOT = Path(__file__).resolve().parents[1]
LOGO = ROOT / "logo" / "kalypsis-logo.jpg"
OUT = ROOT / "docs" / "Kalypsis_Platform_Greek.docx"

NAVY = RGBColor(0x0B, 0x25, 0x45)
AMBER = RGBColor(0xF6, 0xA6, 0x23)
GREY = RGBColor(0x60, 0x60, 0x60)


def set_cell_bg(cell, hex_color: str):
    tc_pr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tc_pr.append(shd)


def add_heading(doc, text, level=1, color=NAVY):
    p = doc.add_paragraph()
    run = p.add_run(text)
    run.bold = True
    run.font.color.rgb = color
    if level == 1:
        run.font.size = Pt(22)
        p.paragraph_format.space_before = Pt(18)
        p.paragraph_format.space_after = Pt(8)
    elif level == 2:
        run.font.size = Pt(15)
        p.paragraph_format.space_before = Pt(12)
        p.paragraph_format.space_after = Pt(4)
    else:
        run.font.size = Pt(12)
        p.paragraph_format.space_before = Pt(8)
        p.paragraph_format.space_after = Pt(2)
    return p


def add_para(doc, text, *, bold=False, italic=False, size=11, color=None, align=None,
             space_after=6):
    p = doc.add_paragraph()
    if align:
        p.alignment = align
    run = p.add_run(text)
    run.bold = bold
    run.italic = italic
    run.font.size = Pt(size)
    if color is not None:
        run.font.color.rgb = color
    p.paragraph_format.space_after = Pt(space_after)
    return p


def add_bullets(doc, items):
    for item in items:
        p = doc.add_paragraph(style="List Bullet")
        if isinstance(item, tuple):
            title, body = item
            r1 = p.add_run(title)
            r1.bold = True
            r1.font.size = Pt(11)
            r2 = p.add_run(" — " + body)
            r2.font.size = Pt(11)
        else:
            r = p.add_run(item)
            r.font.size = Pt(11)
        p.paragraph_format.space_after = Pt(2)


def add_kpi_row(doc, items):
    table = doc.add_table(rows=1, cols=len(items))
    table.autofit = True
    for cell, (number, label) in zip(table.rows[0].cells, items):
        set_cell_bg(cell, "0B2545")
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        p1 = cell.paragraphs[0]
        p1.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r1 = p1.add_run(number)
        r1.bold = True
        r1.font.size = Pt(20)
        r1.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        p2 = cell.add_paragraph()
        p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r2 = p2.add_run(label)
        r2.font.size = Pt(9)
        r2.font.color.rgb = RGBColor(0xE0, 0xE6, 0xEE)
    doc.add_paragraph()


def add_callout(doc, title, body, bg="F4F7FB"):
    table = doc.add_table(rows=1, cols=1)
    cell = table.rows[0].cells[0]
    set_cell_bg(cell, bg)
    p1 = cell.paragraphs[0]
    r1 = p1.add_run(title)
    r1.bold = True
    r1.font.color.rgb = NAVY
    r1.font.size = Pt(12)
    p2 = cell.add_paragraph()
    r2 = p2.add_run(body)
    r2.font.size = Pt(11)
    p2.paragraph_format.space_after = Pt(0)
    doc.add_paragraph()


def add_module_table(doc, rows):
    table = doc.add_table(rows=1 + len(rows), cols=2)
    table.style = "Light Grid"
    header = table.rows[0]
    header.cells[0].text = "Ενότητα"
    header.cells[1].text = "Τι περιλαμβάνει"
    for cell in header.cells:
        set_cell_bg(cell, "0B2545")
        for p in cell.paragraphs:
            for r in p.runs:
                r.bold = True
                r.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
                r.font.size = Pt(10)
    for idx, (name, what) in enumerate(rows, start=1):
        row = table.rows[idx]
        row.cells[0].text = ""
        r = row.cells[0].paragraphs[0].add_run(name)
        r.bold = True
        r.font.size = Pt(10)
        r.font.color.rgb = NAVY
        row.cells[1].text = ""
        r2 = row.cells[1].paragraphs[0].add_run(what)
        r2.font.size = Pt(10)
    doc.add_paragraph()


def add_page_break(doc):
    doc.add_page_break()


def main():
    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Cm(1.8)
        section.bottom_margin = Cm(1.8)
        section.left_margin = Cm(2.2)
        section.right_margin = Cm(2.2)

    # ---------- COVER ----------
    if LOGO.exists():
        p = doc.add_paragraph()
        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        run = p.add_run()
        run.add_picture(str(LOGO), width=Inches(2.2))

    add_para(doc, "KALYPSIS", bold=True, size=42, color=NAVY,
             align=WD_ALIGN_PARAGRAPH.CENTER, space_after=2)
    add_para(doc, "Η Νέα Γενιά Πλατφόρμας Διαχείρισης Ασφαλιστικών Γραφείων",
             italic=True, size=14, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=24)

    add_para(doc,
             "Παρουσίαση Πλατφόρμας — Για Ασφαλιστικά Γραφεία, Μεσίτες & Συνεργαζόμενα Δίκτυα",
             bold=True, size=12, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, space_after=8)
    add_para(doc, "Ελληνική αγορά · Πολυενοικιακό SaaS · Cloud-first",
             italic=True, size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER,
             space_after=32)

    add_kpi_row(doc, [
        ("22+", "Ενεργές ενότητες"),
        ("9", "Επίπεδα προμηθειών"),
        ("100%", "Cloud · Mobile-ready"),
        ("EL/EN", "Δίγλωσσο UI"),
    ])

    add_para(doc,
             "Το παρόν έγγραφο περιγράφει αναλυτικά τι είναι το Kalypsis, "
             "τι προσφέρει στο γραφείο σας, ποιες λειτουργίες ενοποιεί και ποιο είναι το "
             "οδικός χάρτης ανάπτυξης. Σας ευχαριστούμε για το ενδιαφέρον σας.",
             size=11, color=GREY, align=WD_ALIGN_PARAGRAPH.CENTER)

    add_page_break(doc)

    # ---------- 1. Τι είναι το KALYPSIS ----------
    add_heading(doc, "1. Τι είναι το Kalypsis", level=1)
    add_para(doc,
             "Το Kalypsis είναι μια ολοκληρωμένη, ελληνική-first πλατφόρμα διαχείρισης "
             "ασφαλιστικών γραφείων που ενοποιεί σε ένα σύστημα όλες τις λειτουργίες που "
             "παραδοσιακά απαιτούσαν πολλά διαφορετικά εργαλεία: το CRM πελατών, την έκδοση "
             "και παρακολούθηση συμβολαίων, την αποστολή και είσπραξη ασφαλίστρων, τον "
             "υπολογισμό προμηθειών και υπερπρομηθειών, τη λογιστική, το marketing και την "
             "επικοινωνία με ασφαλιστικές εταιρείες, τράπεζες και τρίτες εφαρμογές.")
    add_para(doc,
             "Σχεδιάστηκε με γνώμονα τις πραγματικές απαιτήσεις της ελληνικής αγοράς — "
             "ΑΦΜ, ΚΕΠΥΟ, ΔΙΑΣ, ελληνικά κείμενα παντού — και υλοποιείται με σύγχρονη "
             "τεχνολογία ώστε να τρέχει εξίσου καλά σε desktop, tablet και κινητό.")

    add_callout(doc,
                "Η φιλοσοφία μας",
                "Ένα γραφείο, ένα σύστημα. Όχι τέσσερα διαφορετικά λογισμικά για να "
                "βγάλετε άκρη με τα συμβόλαια, τις προμήθειες, τα emails και το λογιστή. "
                "Όλα ζωντανά συνδεδεμένα, διαθέσιμα από οπουδήποτε, για όλους τους ρόλους.")

    add_heading(doc, "Σε ποιους απευθύνεται", level=2)
    add_bullets(doc, [
        ("Ασφαλιστικά γραφεία", "Από μικρά εξειδικευμένα γραφεία μέχρι δίκτυα με δεκάδες συνεργάτες."),
        ("Μεσίτες ασφαλίσεων", "Διαχείριση πολλαπλών εταιρειών και χαρτοφυλακίων σε ένα σημείο."),
        ("Ανεξάρτητοι διαμεσολαβητές", "Lite plans για έναν χρήστη με πλήρες παραγωγικό circuit."),
        ("Συνεργατικά δίκτυα", "Πυραμίδες έως 9 επιπέδων με αυτόματο υπολογισμό υπερπρομηθειών."),
    ])

    # ---------- 2. Οι ρόλοι ----------
    add_heading(doc, "2. Οι Ρόλοι της Πλατφόρμας", level=1)
    add_para(doc,
             "Το Kalypsis ξεχωρίζει έξι αυτόνομους ρόλους, καθένας με δικό του dashboard "
             "και αυστηρό σύστημα δικαιωμάτων:")
    add_bullets(doc, [
        ("Διαχειριστής Πλατφόρμας (Super Admin)",
         "Εποπτεύει όλα τα γραφεία της πλατφόρμας, εγκρίνει εγγραφές, παρακολουθεί υγεία "
         "συστήματος και μπορεί να «μπει ως» οποιοδήποτε γραφείο για υποστήριξη."),
        ("Υπάλληλος Πλατφόρμας",
         "Υποστήριξη πελατών και τεχνική παρακολούθηση χωρίς διαχειριστικά δικαιώματα."),
        ("Διαχειριστής Γραφείου (Agency Admin)",
         "Πλήρης έλεγχος του γραφείου του: πελάτες, συμβόλαια, οικονομικά, προμήθειες, "
         "υπάλληλοι, marketing και ρυθμίσεις."),
        ("Υπάλληλος Γραφείου (Agency User)",
         "Καθημερινή χρήση χωρίς πρόσβαση σε ευαίσθητες ρυθμίσεις, API tokens ή προμήθειες."),
        ("Συνεργάτης / Διαμεσολαβητής (Producer)",
         "Δικό του portal: τα συμβόλαια που έφερε, οι προμήθειές του, οι πελάτες του."),
        ("Πελάτης (Customer)",
         "Δικό του portal: συμβόλαια σε PDF, εκτύπωση, αιτήματα, ειδοποιήσεις λήξης."),
    ])

    add_callout(doc, "Mode «Έλεγχος Αστυνομίας» για πελάτες",
                "Ο πελάτης μπορεί από το κινητό του να εμφανίσει σε fullscreen mode τα "
                "στοιχεία του ενεργού συμβολαίου του (αριθμός, εταιρεία, διάρκεια) με "
                "εμφανές «πιστοποιημένο» σήμα — ιδανικό για έλεγχο στο δρόμο.")

    # ---------- 3. Τι Προσφέρει στο Γραφείο ----------
    add_heading(doc, "3. Τι Προσφέρει στο Γραφείο σας", level=1)

    add_heading(doc, "3.1 Διαχείριση Πελατών (CRM)", level=2)
    add_bullets(doc, [
        "Πλήρης καρτέλα πελάτη (φυσικά πρόσωπα και επιχειρήσεις) με αυτόματη αρίθμηση.",
        "Αυτόματη δημιουργία λογαριασμού πελάτη στο portal με προσωρινό κωδικό.",
        "Ιστορικό συμβολαίων, εγγράφων, εισπράξεων και επικοινωνίας σε ένα σημείο.",
        "Έξυπνη αναζήτηση: όνομα, ΑΦΜ, email, αριθμός πελάτη.",
    ])

    add_heading(doc, "3.2 Συμβόλαια & Cover Notes", level=2)
    add_bullets(doc, [
        "Έκδοση συμβολαίων με 7 προκαθορισμένους κλάδους και custom πεδία.",
        "Cover Notes: άμεσα προσωρινά πιστοποιητικά με εκτύπωση/PDF.",
        "Παρακολούθηση λήξεων, ανανεώσεις και αυτόματες ειδοποιήσεις.",
        "Upload εγγράφων (PDF, JPG) ανά συμβόλαιο, ορατά από τον πελάτη.",
        "Παρακολούθηση παράδοσης (email, courier, αυτοπροσώπως) με αυτόματη ταξινόμηση.",
    ])

    add_heading(doc, "3.3 Cl­ient Portal (B2C)", level=2)
    add_bullets(doc, [
        "Branded portal για κάθε πελάτη του γραφείου σας με το δικό σας λογότυπο.",
        "Λήψη, εκτύπωση και προβολή συμβολαίων από οποιοδήποτε device.",
        "Mobile-first: ο πελάτης βλέπει τα συμβόλαια από το κινητό σαν app.",
        "Άμεσες ειδοποιήσεις για λήξεις, νέα έγγραφα και αιτήματα.",
        "Wizard νέας ασφάλισης σε 3 βήματα που στέλνεται κατευθείαν στο γραφείο σας.",
    ])

    add_heading(doc, "3.4 Συνεργατικό Δίκτυο", level=2)
    add_bullets(doc, [
        "Καταχώρηση συνεργατών (Producers) με κατάσταση Active/Inactive/Terminated.",
        "Έκδοση πρόσβασης στο B2B portal με ένα κλικ (αυτόματος κωδικός).",
        "Δικαιώματα ανά συνεργάτη: έκδοση συμβολαίων, προβολή προμηθειών, πελάτες.",
        "Πυραμίδα έως 9 επιπέδων υπερπρομηθειών με δικά σας ποσοστά.",
    ])

    add_heading(doc, "3.5 Παραμετρικοί Τιμοκατάλογοι (Tariffs)", level=2)
    add_bullets(doc, [
        "Πολλαπλοί τιμοκατάλογοι ανά εταιρεία × κλάδο × αντικείμενο ασφάλισης.",
        "Συντελεστές (ηλικίας, ζώνης, ισχύος) με JSON ή UI editor.",
        "Έναρξη/λήξη ισχύος για ιστορικότητα και αυτόματη επιλογή ταρίφας.",
        "Σύνδεση κατευθείαν με την έκδοση συμβολαίου ή Cover Note.",
    ])

    add_heading(doc, "3.6 Σχεδιαστής Κλάδων (Branch Designer)", level=2)
    add_bullets(doc, [
        "Δημιουργία νέων κλάδων ασφάλισης χωρίς ανάγκη development.",
        "Custom πεδία (text, number, dropdown) ανά κλάδο.",
        "Καλύψεις, πακέτα και επιλογές με συντελεστές προμήθειας.",
        "Δικό σας ερωτηματολόγιο για online έκδοση.",
    ])

    add_page_break(doc)

    add_heading(doc, "3.7 Οικονομικά (Financial Circuits)", level=2)
    add_para(doc,
             "Το Kalypsis ενοποιεί όλο τον οικονομικό κύκλο σε ένα σύστημα. Κάθε είσπραξη "
             "και κάθε πληρωμή δημιουργεί αυτόματα την αντίστοιχη οικονομική κίνηση που "
             "ενημερώνει την καρτέλα πελάτη/συνεργάτη/εταιρείας σε πραγματικό χρόνο.")
    add_bullets(doc, [
        ("Εισπράξεις",
         "Μετρητά, κάρτα, έμβασμα, επιταγή, συναλλαγματική. Αυτόματη απόδειξη με αρίθμηση."),
        ("Πληρωμές με netting προμηθειών",
         "Εξόφληση εταιρειών μείον τις προμήθειες σας, σε μία μόνο εγγραφή."),
        ("Αξιόγραφα",
         "Επιταγές και συναλλαγματικές με χαρτοφυλάκιο λήξεων και κατάσταση εξόφλησης."),
        ("Καρτέλες & Καρτέλα Κίνησης",
         "Live ledger για κάθε πελάτη/συνεργάτη/εταιρεία με φίλτρα και export."),
        ("KPIs σε πραγματικό χρόνο",
         "Εισπράξεις, πληρωμές, προμήθειες ανά μήνα σε ζωντανό dashboard."),
    ])

    add_heading(doc, "3.8 Προμήθειες (Commissions)", level=2)
    add_bullets(doc, [
        "Παραμετρικοί κανόνες ανά Συνεργάτη × Αντικείμενο Ασφάλισης × Κάλυψη.",
        "Διαφορετικός συντελεστής για ετήσιες αναπροσαρμογές ασφαλίστρου.",
        "Splits μεταξύ πολλαπλών συνεργατών στο ίδιο συμβόλαιο.",
        "Πλήρες audit trail κάθε αλλαγής για ιστορικότητα.",
    ])

    add_heading(doc, "3.9 Υπερπρομήθειες (Over-commissions)", level=2)
    add_bullets(doc, [
        "Πυραμίδα έως 9 επίπεδα (manager → senior → agent…).",
        "Ποσοστό ανά επίπεδο και ανά κλάδο.",
        "Ξεχωριστές εκκαθαρίσεις ανά πυραμίδα.",
        "Visualization των ομάδων και των ροών εσόδων.",
    ])

    add_heading(doc, "3.10 Παραγωγή & Στόχοι", level=2)
    add_bullets(doc, [
        "Αναλυτικά στατιστικά ανά μήνα, κλάδο, εταιρεία, συνεργάτη.",
        "Bar/Pie charts με YoY σύγκριση και export σε Excel/PDF.",
        "Στόχοι μηνιαίοι / ετήσιοι ανά συνεργάτη ή για όλο το γραφείο.",
        "Retention rate και πρόοδος σε live progress bars.",
    ])

    add_heading(doc, "3.11 Marketing", level=2)
    add_bullets(doc, [
        "Μαζικές καμπάνιες email σε segments πελατών (όλοι, λήγουν σύντομα, με email).",
        "HTML editor με δικά σας templates.",
        "Auto-recipient calculation με βάση το segment.",
        "Καταγραφή αποστολών στην καρτέλα του κάθε πελάτη.",
    ])

    add_heading(doc, "3.12 Ραντεβού & Ημερολόγιο", level=2)
    add_bullets(doc, [
        "Καταχώρηση ραντεβού με πελάτες/συνεργάτες ανά υπάλληλο.",
        "Ομαδοποίηση κατά ημέρα με συνοπτικά cards.",
        "Σύνδεση ραντεβού με πελάτη και συμβόλαιο.",
        "Καταστάσεις: Προγραμματισμένο / Ολοκληρώθηκε / Ακυρωμένο.",
    ])

    add_heading(doc, "3.13 Document Manager", level=2)
    add_bullets(doc, [
        "Κεντρική διαχείριση εγγράφων ανά πελάτη με φακέλους και χρώματα.",
        "Σύνδεση κάθε εγγράφου με συμβόλαιο και κανόνες πρόσβασης.",
        "Άμεση παράδοση στον πελάτη μέσω portal.",
    ])

    # ---------- 4. Τραπεζικές & Λογιστικές Ολοκληρώσεις ----------
    add_heading(doc, "4. Τραπεζικές & Λογιστικές Ολοκληρώσεις", level=1)

    add_heading(doc, "4.1 ΔΙΑΣ (Διατραπεζική)", level=2)
    add_bullets(doc, [
        "Αυτόματη παραγωγή RF code ανά συμβόλαιο.",
        "Παρακολούθηση κατάστασης πληρωμής (Pending / Paid / Cancelled).",
        "Καταγραφή bank reference στην εξόφληση.",
        "Roadmap: Αυτόματη έκδοση συμβολαίου μόλις φτάσει η επιβεβαίωση τράπεζας.",
    ])

    add_heading(doc, "4.2 Συνδέσεις Τραπεζών", level=2)
    add_bullets(doc, [
        "Καταχώρηση τραπεζικών λογαριασμών (IBAN, BIC, όνομα κατόχου).",
        "Sync timestamp και κατάσταση ενεργοποίησης.",
        "Roadmap: Open Banking sync για αυτόματη συμφωνία καταθέσεων.",
    ])

    add_heading(doc, "4.3 Λογιστική Εξαγωγή", level=2)
    add_bullets(doc, [
        "Παραγωγή αρχείου ανά μήνα με όλες τις οικονομικές κινήσεις.",
        "Ιστορικό runs με αρίθμηση εγγραφών.",
        "Roadmap: Direct connectors σε BlueByte και κορυφαία ελληνικά λογιστικά.",
    ])

    add_heading(doc, "4.4 ΚΕΠΥΟ", level=2)
    add_bullets(doc, [
        "Ετήσια συγκεντρωτική κατάσταση πελατών/προμηθευτών.",
        "Auto-fill από οικονομικές κινήσεις του έτους.",
        "Roadmap: XML για ΑΑΔΕ.",
    ])

    add_heading(doc, "4.5 Εισαγωγή από Μαγνητικά Μέσα", level=2)
    add_bullets(doc, [
        "Upload CSV/XML αρχείων ασφαλιστικών εταιρειών.",
        "Auto-match παραγωγής σε συμβόλαια.",
        "Αναφορά γραμμών, matches και σφαλμάτων.",
    ])

    # ---------- 5. Portals & API ----------
    add_heading(doc, "5. Portals για Συνεργάτες & Τρίτες Εφαρμογές", level=1)
    add_bullets(doc, [
        ("B2B Portal Συνεργατών",
         "Λογαριασμοί πρόσβασης ανά Producer με granular δικαιώματα."),
        ("Client Portal",
         "Branded πύλη με το λογότυπο του γραφείου για κάθε πελάτη."),
        ("Third-party API",
         "Έκδοση API keys με prefix + SHA-256 hash. Scopes ανά κλειδί. Token εμφανίζεται "
         "μόνο μία φορά κατά την έκδοση για ασφάλεια."),
    ])

    add_page_break(doc)

    # ---------- 6. Branding ----------
    add_heading(doc, "6. Branding & Εξατομίκευση", level=1)
    add_bullets(doc, [
        "Upload λογότυπου γραφείου (PNG/JPEG/SVG/WebP έως 4 MB).",
        "Το λογότυπο εμφανίζεται σε όλα τα dashboards του γραφείου και του πελάτη.",
        "Χρώμα brand, στοιχεία επικοινωνίας, ΑΦΜ, διεύθυνση, νόμισμα.",
        "Διάρκεια συμβολαίου default (μήνες) για ταχύτερη εισαγωγή.",
    ])

    # ---------- 7. Ασφάλεια & Συμμόρφωση ----------
    add_heading(doc, "7. Ασφάλεια & Συμμόρφωση", level=1)
    add_bullets(doc, [
        ("Multi-tenant με αυστηρή απομόνωση",
         "Κάθε γραφείο βλέπει μόνο τα δικά του δεδομένα — επιβεβαιωμένο σε επίπεδο βάσης."),
        ("JWT + refresh tokens",
         "Σύντομα access tokens, refresh tokens που μπορούν να ανακληθούν."),
        ("Bcrypt για κωδικούς, SHA-256 για tokens",
         "Πιστοποίηση industry-standard, χωρίς αποθήκευση plaintext."),
        ("Audit Log",
         "Κάθε δημιουργία/αλλαγή/διαγραφή καταγράφεται με χρήστη, χρόνο και old→new values."),
        ("Soft delete παντού",
         "Καμία οριστική απώλεια δεδομένων χωρίς explicit action."),
        ("Role-based authorization",
         "5 policies (PlatformAdmin, PlatformLevel, AgencyAdmin, AgencyStaff, Producer)."),
        ("GDPR-friendly",
         "Προφίλ πελάτη, δικαίωμα διαγραφής, καταγραφή συγκαταθέσεων (cookies banner)."),
    ])

    # ---------- 8. Εμπειρία Χρήστη ----------
    add_heading(doc, "8. Εμπειρία Χρήστη", level=1)
    add_bullets(doc, [
        ("Δίγλωσσο UI",
         "Ελληνικά πρώτη γλώσσα, αγγλικά δεύτερη. Toggler με σημαίες σε κάθε σελίδα."),
        ("Mobile-first",
         "Όλα τα dashboards δουλεύουν άψογα σε desktop, tablet και κινητό. "
         "Drawer overlay στο mobile, persistent στο desktop."),
        ("Εσωτερικός scroll σε πίνακες",
         "Δεν χάνεται η πλοήγηση όταν ένας πίνακας είναι πλατύς."),
        ("«Να με θυμάσαι»",
         "Επιλογή για διατήρηση session με localStorage vs sessionStorage."),
        ("Forgot password με email",
         "Αυτόματη αποστολή reset link μέσω Brevo με το branding του γραφείου."),
    ])

    # ---------- 9. Διαχείριση Πλατφόρμας (για Admin Πλατφόρμας) ----------
    add_heading(doc, "9. Διαχείριση Πλατφόρμας (για Platform Admin)", level=1)
    add_bullets(doc, [
        ("Επισκόπηση Γραφείων",
         "Stat tiles ανά γραφείο: χρήστες, πελάτες, συμβόλαια, ασφάλιστρα, ζημιές."),
        ("«Είσοδος ως Γραφείο»",
         "Ο Super Admin μπορεί με ένα κλικ να μπει στο dashboard οποιουδήποτε γραφείου "
         "και να βλέπει/ενεργεί σαν AgencyAdmin για υποστήριξη."),
        ("Cross-tenant CRUD",
         "Διαχείριση όλων των χρηστών της πλατφόρμας με φίλτρα και αναζήτηση."),
        ("Bulk Actions",
         "Μαζική ενεργοποίηση, απενεργοποίηση, διαγραφή χρηστών."),
        ("Audit Log Πλατφόρμας",
         "Όλες οι ενέργειες σε όλα τα γραφεία σε ένα σημείο."),
        ("Σύντομα",
         "Carriers, subscription plans, billing, email templates, broadcast, branding, "
         "integrations, backups, storage, jobs, status, compliance, support."),
    ])

    add_page_break(doc)

    # ---------- 10. Πίνακας Ενοτήτων ----------
    add_heading(doc, "10. Συνολικός Πίνακας Ενοτήτων (Ζωντανές & Roadmap)", level=1)
    add_module_table(doc, [
        ("Πελάτες", "CRM, καρτέλες, portal accounts, αναζήτηση"),
        ("Συμβόλαια", "Έκδοση, ανανέωση, ακύρωση, έγγραφα, παραστατικά"),
        ("Cover Notes", "Προσωρινά πιστοποιητικά με εκτύπωση"),
        ("Τιμοκατάλογοι", "Παραμετρικά pricing per branch × company"),
        ("Σχεδιαστής Κλάδων", "Custom fields, coverages, ερωτηματολόγια"),
        ("Παράδοση Συμβολαίων", "Tracking ανά κανάλι (email/courier/in-person)"),
        ("Ραντεβού", "Calendar με σύνδεση σε πελάτη και συμβόλαιο"),
        ("Document Manager", "Φάκελοι ανά πελάτη με χρωματική κωδικοποίηση"),
        ("Marketing", "Καμπάνιες email με smart segments"),
        ("B2B / B2C Portals", "Συνεργάτες & πελάτες με branded πρόσβαση"),
        ("Προμήθειες", "Παραμετρικοί κανόνες με ιστορικότητα"),
        ("Υπερπρομήθειες", "9-επίπεδη πυραμίδα"),
        ("Στατιστικά Παραγωγής", "Charts, YoY, ανά συνεργάτη/κλάδο"),
        ("Στόχοι & Retention", "Live progress, bonuses"),
        ("Οικονομικές Κινήσεις", "Ledger με filters και summary"),
        ("Εισπράξεις", "Πολλαπλοί τρόποι πληρωμής, αυτόματη απόδειξη"),
        ("Πληρωμές", "Netting προμηθειών"),
        ("Αξιόγραφα", "Επιταγές/συναλλαγματικές με ημερολόγιο λήξεων"),
        ("ΔΙΑΣ", "RF codes, mark-paid flow"),
        ("Συνδέσεις Τραπεζών", "IBAN/BIC + sync timestamp"),
        ("Third-party API", "Tokens με scopes, hash, expiration"),
        ("Λογιστική", "Εξαγωγή ανά μήνα"),
        ("ΚΕΠΥΟ", "Ετήσιες συγκεντρωτικές"),
        ("Μαγνητικά Μέσα", "Import παραγωγής από αρχεία ασφαλιστικών"),
        ("Audit Log", "Παντού, με old/new values"),
        ("Branding", "Logo, χρώματα, στοιχεία επικοινωνίας"),
    ])

    # ---------- 11. Roadmap ----------
    add_heading(doc, "11. Roadmap — Τι Έρχεται", level=1)
    add_bullets(doc, [
        "Open Banking integration για αυτόματη συμφωνία τράπεζας.",
        "Direct connectors προς ελληνικά λογιστικά συστήματα.",
        "Αυτόματος υπολογισμός ποινών μη-ανανέωσης και suggested actions.",
        "AI-assisted κατηγοριοποίηση εγγράφων και OCR με full-text search.",
        "Παραμετρικές καμπάνιες WhatsApp/Viber.",
        "Mobile native apps για πελάτες και συνεργάτες.",
        "Επέκταση φόρμας ΚΕΠΥΟ στην επίσημη XML ΑΑΔΕ.",
        "Marketplace από πρόσθετες εφαρμογές και integrations.",
    ])

    # ---------- 12. Τεχνολογία ----------
    add_heading(doc, "12. Τεχνολογία & Αρχιτεκτονική", level=1)
    add_bullets(doc, [
        (".NET 10 + Clean Architecture",
         "Διαχωρισμός σε Domain / Application / Infrastructure / API."),
        ("React 18 + Vite 5 + TypeScript",
         "Σύγχρονο, γρήγορο UI με MUI 6 και TanStack Query."),
        ("MySQL (Pomelo) + EF Core 9",
         "Σταθερή βάση δεδομένων με αυτόματες migrations."),
        ("MediatR CQRS + FluentValidation",
         "Καθαρή ροή εντολών και ερωτημάτων με validation pipeline."),
        ("JWT auth + Refresh tokens",
         "Πολυχρηστικό multi-tenant με policy-based authorization."),
        ("Brevo για email",
         "Transactional emails (welcome, reset password, ειδοποιήσεις)."),
        ("Local file storage σε dev / S3-compatible σε production",
         "Έγγραφα και λογότυπα με abstracted IFileStorage."),
        ("CI-ready",
         "Όλα ελεγμένα με ενσωματωμένα tests, audit logs και observability hooks."),
    ])

    # ---------- 13. Έναρξη Συνεργασίας ----------
    add_heading(doc, "13. Πώς Ξεκινάμε", level=1)
    add_para(doc,
             "Η μετάβαση στο Kalypsis γίνεται με ελάχιστο χρόνο και χωρίς διακοπή λειτουργίας:")
    add_bullets(doc, [
        ("1. Δημιουργία γραφείου σας στην πλατφόρμα",
         "Σε λιγότερο από 5 λεπτά. Παίρνετε άμεσα admin λογαριασμό."),
        ("2. Εισαγωγή πελατών και συνεργατών",
         "Είτε χειροκίνητα είτε μέσω Excel/CSV import."),
        ("3. Εκπαίδευση",
         "Online session για τους χρήστες σας, γραπτό υλικό στα ελληνικά."),
        ("4. Σταδιακή μετάπτωση παλιών συμβολαίων",
         "Με τη βοήθεια της ομάδας υποστήριξης Kalypsis."),
        ("5. Go-live",
         "Από την πρώτη μέρα μπορείτε να εκδίδετε συμβόλαια και Cover Notes."),
    ])

    add_callout(doc, "Εγγύηση επιστροφής 30 ημερών",
                "Δοκιμάστε όλες τις λειτουργίες της πλατφόρμας. Αν για οποιονδήποτε λόγο "
                "δεν σας καλύπτει, σας επιστρέφουμε ολόκληρο το ποσό χωρίς ερωτήσεις.")

    # ---------- 14. Επαφή ----------
    add_heading(doc, "14. Στοιχεία Επικοινωνίας", level=1)
    add_para(doc, "Επικοινωνήστε μαζί μας για demo, ερωτήσεις ή προσφορά:", size=11)
    add_bullets(doc, [
        "Email: hello@kalypsis.gr",
        "Τηλέφωνο: +30 210 000 0000",
        "Web: www.kalypsis.gr",
        "Github: github.com/PanagiotisKotsorgios/kalypsis",
    ])

    add_para(doc, "", space_after=24)
    add_para(doc, "Σας ευχαριστούμε για το χρόνο σας.",
             bold=True, color=NAVY, align=WD_ALIGN_PARAGRAPH.CENTER, size=12)
    add_para(doc, "— Η ομάδα του Kalypsis", italic=True, color=GREY,
             align=WD_ALIGN_PARAGRAPH.CENTER, size=11)

    OUT.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT))
    print(f"OK · wrote {OUT}")


if __name__ == "__main__":
    main()
